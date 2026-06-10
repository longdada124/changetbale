import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import re
from datetime import datetime, timedelta

st.set_page_config(page_title="課表彙整與代調課系統", layout="wide")

# --- 功能 1：簡易密碼保護機制 ---
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.markdown("<h2 style='text-align: center;'>🔐 課表彙整系統 - 系統登入</h2>", unsafe_allow_html=True)
    col_l, col_m, col_r = st.columns([1, 1.5, 1])
    with col_m:
        password = st.text_input("請輸入系統驗證密碼：", type="password")
        if st.button("確認登入", use_container_width=True):
            if password == "1030018":
                st.session_state.authenticated = True
                st.success("🎉 登入成功！")
                st.rerun()
            else:
                st.error("❌ 密碼錯誤，請重新輸入。")
    st.stop()

# --- 核心工具函數：產生週次日期區間 ---
def generate_weeks(start_date_str, total_weeks=20):
    start_date = datetime.strptime(start_date_str, "%Y-%m-%d")
    # 調整至該週的週一
    start_date = start_date - timedelta(days=start_date.weekday())
    weeks_info = {}
    day_names = ["一", "二", "三", "四", "五", "六", "日"]
    
    for w in range(1, total_weeks + 1):
        w_start = start_date + timedelta(weeks=w-1)
        w_end = w_start + timedelta(days=4) # 週一到週五
        label = f"第{w:02d}週 ({w_start.strftime('%Y-%m-%d')} ~ {w_end.strftime('%Y-%m-%d')})"
        
        # 建立這一週每一天對應的日期與星期對照
        days_map = {}
        date_to_day_num = {}
        for idx in range(5):
            current_day = w_start + timedelta(days=idx)
            days_map[day_names[idx]] = current_day.strftime("%Y-%m-%d")
            days_map[f"週{day_names[idx]}"] = current_day.strftime("%Y-%m-%d")
            date_to_day_num[current_day.strftime("%Y-%m-%d")] = (idx + 1, day_names[idx]) # (1~5, 星期)
            
        weeks_info[w] = {
            "label": label,
            "start": w_start,
            "end": w_end,
            "days": days_map,
            "date_to_day_num": date_to_day_num
        }
    return weeks_info

# --- 核心替換函數 ---
def master_replace(doc_obj, old_text, new_text):
    if isinstance(new_text, (float, int)):
        new_val = str(int(new_text))
    else:
        new_val = str(new_text) if (new_text and str(new_text).strip() != "") else ""
    targets = list(doc_obj.paragraphs)
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                targets.extend(cell.paragraphs)
    for p in targets:
        if old_text in p.text:
            full_text = "".join([run.text for run in p.runs])
            updated_text = full_text.replace(old_text, new_val)
            for i, run in enumerate(p.runs):
                run.text = updated_text if i == 0 else ""

def load_default_template(file_name):
    try:
        with open(file_name, "rb") as f:
            return f.read()
    except FileNotFoundError:
        return None

# --- 側邊欄：資料管理 ---
with st.sidebar:
    st.header("⚙️ 資料管理")
    
    if st.button("🔒 安全登出系統", type="secondary", use_container_width=True):
        st.session_state.authenticated = False
        st.rerun()
        
    st.divider()
    if st.button("🧹 清空重置系統", use_container_width=True):
        for key in list(st.session_state.keys()): 
            if key != "authenticated": del st.session_state[key]
        st.rerun()

    st.divider()
    st.subheader("📅 學期時間設定")
    # 動態調整學期基準日，預設以當前年度春季學期為例
    semester_start = st.date_input("設定本學期【第一週週一】日期", datetime(2026, 2, 16))
    
    st.divider()
    st.subheader("📥 範本下載")
    data_templates = {
        "1. 配課表範本": "配課表.xlsx",
        "2. 課表範本": "課表.xlsx",
        "3. 教師排序表範本": "教師排序表.xlsx"
    }
    for label, file_name in data_templates.items():
        try:
            with open(file_name, "rb") as f:
                st.download_button(label=f"{label}", data=f, file_name=file_name, key=f"dl_{file_name}")
        except FileNotFoundError:
            st.caption(f"⚠️ 找不到 {file_name}")
            
    st.divider()
    st.subheader("📤 上傳資料檔")
    f_assign = st.file_uploader("1. 上傳【配課表】", type=["xlsx", "csv"])
    f_time = st.file_uploader("2. 上傳【課表】", type=["xlsx", "csv"])
    f_sort = st.file_uploader("3. 上傳【教師排序暨時數表】", type=["xlsx", "csv"])
    
    if f_assign and f_time and st.button("🚀 執行整合", use_container_width=True):
        class_temp = load_default_template("班級樣板.docx")
        teacher_temp = load_default_template("教師樣板.docx")
        
        if not class_temp or not teacher_temp:
            st.error("❌ 後台找不到樣板，請確認檔案。")
        else:
            with st.spinner("系統重構與多週次課表同步中..."):
                df_assign = pd.read_csv(f_assign) if f_assign.name.endswith('.csv') else pd.read_excel(f_assign)
                df_time = pd.read_csv(f_time) if f_time.name.endswith('.csv') else pd.read_excel(f_time)
                
                # 初始化 20 週時間軸
                weeks_db = generate_weeks(semester_start.strftime("%Y-%m-%d"), total_weeks=20)
                st.session_state.weeks_db = weeks_db
                st.session_state.class_template = class_temp
                st.session_state.teacher_template = teacher_temp

                # 1. 解析配課
                assign_lookup, all_teachers_db, tutors = [], set(), {}
                for _, row in df_assign.iterrows():
                    c, s, t_raw = str(row['班級']).strip(), str(row['科目']).strip(), str(row['教師']).strip()
                    t_list = [name.strip() for name in t_raw.split('/')]
                    for t in t_list:
                        if t and t != "nan":
                            assign_lookup.append({'c': c, 's': s, 't': t})
                            all_teachers_db.add(t)
                    if s == "班級": tutors[c] = t_raw

                # 2. 教師排序與基礎時數
                ordered_teachers, base_hours, all_teachers_list = [], {}, list(all_teachers_db)
                if f_sort:
                    df_s = pd.read_csv(f_sort) if f_sort.name.endswith('.csv') else pd.read_excel(f_sort)
                    for _, s_row in df_s.iterrows():
                        t_name = str(s_row.iloc[0]).strip()
                        if t_name in all_teachers_list:
                            ordered_teachers.append(t_name)
                            try: base_hours[t_name] = int(s_row.iloc[1])
                            except: base_hours[t_name] = 0
                    for t in all_teachers_list:
                        if t not in ordered_teachers: ordered_teachers.append(t); base_hours[t] = 0
                else:
                    ordered_teachers = sorted(all_teachers_list)
                    base_hours = {t: 0 for t in ordered_teachers}

                # 3. 解析原始基礎課表 (不分週次)
                base_class_data, base_teacher_data = {}, {}
                day_map = {"一":1,"二":2,"三":3,"四":4,"五":5,"週一":1,"週二":2,"週三":3,"週四":4,"週五":5}
                
                for _, row in df_time.iterrows():
                    c_raw, s_raw = str(row['班級']).strip(), str(row['科目']).strip()
                    d, p_match = day_map.get(str(row['星期']).strip(), 0), re.search(r'\d+', str(row['節次']))
                    if not (p_match and d > 0): continue
                    p = int(p_match.group())

                    if not s_raw or s_raw == "nan" or s_raw == "":
                        display_t = ""
                        s_raw = ""
                        curr_t_list = []
                    else:
                        curr_t_list = [item['t'] for item in assign_lookup if item['c'] == c_raw and item['s'] == s_raw]
                        display_t = "/".join(curr_t_list) if curr_t_list else "未知教師"
                    
                    if c_raw not in base_class_data: base_class_data[c_raw] = {}
                    base_class_data[c_raw][(d, p)] = {"subj": s_raw, "teacher": display_t, "status": "正常", "note": ""}
                    
                    for t in curr_t_list:
                        if t not in base_teacher_data: base_teacher_data[t] = {}
                        base_teacher_data[t][(d, p)] = {"subj": s_raw, "class": c_raw, "status": "正常", "note": ""}

                # 4. 將基礎課表複製延伸到全學期 20 週中，建立獨立資料節點
                schedule_by_week = {}
                for w in range(1, 21):
                    schedule_by_week[w] = {
                        "class_data": {},
                        "teacher_data": {}
                    }
                    # 複製班級
                    for c_key, c_val in base_class_data.items():
                        schedule_by_week[w]["class_data"][c_key] = {k: v.copy() for k, v in c_val.items()}
                    # 複製教師
                    for t_key, t_val in base_teacher_data.items():
                        schedule_by_week[w]["teacher_data"][t_key] = {k: v.copy() for k, v in t_val.items()}

                st.session_state.update({
                    "schedule_by_week": schedule_by_week, 
                    "tutors_map": tutors,
                    "base_hours": base_hours, 
                    "ordered_teachers": ordered_teachers,
                    "sel_class": sorted(list(base_class_data.keys()))[0] if base_class_data else "", 
                    "sel_teacher": ordered_teachers[0] if ordered_teachers else ""
                })
                st.rerun()

# --- 主介面與預覽 ---
if 'schedule_by_week' in st.session_state:
    
    # 頂部篩選列（複刻圖片效果）
    st.markdown("### 🗓️ 全校課表與調度管理中心")
    tc1, tc2, tc3 = st.columns([1, 1.5, 1.5])
    with tc1:
        st.selectbox("學年學期", ["114學年度 下學期"], index=0, disabled=True)
    with tc2:
        week_options = {w: info["label"] for w, info in st.session_state.weeks_db.items()}
        current_w = st.selectbox("選取週次", options=list(week_options.keys()), format_func=lambda x: week_options[x], index=14) # 預設第15週
    
    tab1, tab2, tab3 = st.tabs(["🏫 班級課表預覽", "👩‍🏫 教師課表預覽", "🔄 增減 / 代調課作業"])

    # 動態計算當週全校老師的時數（因為可能會有代調課影響）
    # 初始化當週統計
    current_total_counts = {t: 0 for t in st.session_state.ordered_teachers}
    for t in st.session_state.ordered_teachers:
        t_lessons = st.session_state.schedule_by_week[current_w]["teacher_data"].get(t, {})
        for (d, p), info in t_lessons.items():
            if info["subj"] != "":
                current_total_counts[t] += 1

    # --- Tab 1: 班級課表 ---
    with tab1:
        classes = sorted(list(st.session_state.schedule_by_week[current_w]["class_data"].keys()))
        if classes:
            curr_c = st.session_state.get('sel_class', classes[0])
            if curr_c not in classes: curr_c = classes[0]
            
            b_col1, b_col2, b_col3 = st.columns([1, 2, 1])
            if b_col1.button("⬅️ 上一班", key="prev_c"):
                st.session_state.sel_class = classes[(classes.index(curr_c) - 1) % len(classes)]; st.rerun()
            if b_col3.button("下一班 ➡️", key="next_c"):
                st.session_state.sel_class = classes[(classes.index(curr_c) + 1) % len(classes)]; st.rerun()
            with b_col2: 
                st.session_state.sel_class = st.selectbox("選取班級", classes, index=classes.index(curr_c), key="sb_c")
            
            target_c = st.session_state.sel_class
            st.info(f"📍 班級：{target_c} | 導師：{st.session_state.tutors_map.get(target_c, '未設定')} | 當週：{week_options[current_w]}")
            
            # 建立表格呈現 (高亮代調課)
            c_preview = []
            current_week_days = st.session_state.weeks_db[current_w]["days"]
            
            for p in range(1, 9):
                row = {"節次": f"第 {p} 節"}
                for d_idx, d_name in enumerate(["一", "二", "三", "四", "五"]):
                    d = d_idx + 1
                    info = st.session_state.schedule_by_week[current_w]["class_data"][target_c].get((d,p))
                    if info and info["subj"] != "":
                        status_tag = f" [{info['status']}]" if info['status'] != "正常" else ""
                        row[f"週{d_name}"] = f"{info['subj']}\n👤{info['teacher']}{status_tag}"
                    else:
                        row[f"週{d_name}"] = ""
                c_preview.append(row)
                
            # 使用 dataframe 搭配 st.dataframe 呈現，以便未來加入顏色樣式，或直接用傳統 st.table
            df_display = pd.DataFrame(c_preview)
            st.table(df_display)

            # 下載通知單
            if st.button(f"📥 下載 {target_c} 當週課表樣板檔"):
                doc = Document(BytesIO(st.session_state.class_template))
                master_replace(doc, "{{CLASS}}", target_c)
                master_replace(doc, "{{TUTOR}}", st.session_state.tutors_map.get(target_c, "未設定")) 
                for d, p in [(d,p) for d in range(1,6) for p in range(1,9)]:
                    v = st.session_state.schedule_by_week[current_w]["class_data"][target_c].get((d,p), {"subj":"","teacher":""})
                    status_tag = f"({v.get('status','')})" if v.get('status','正常') != '正常' else ""
                    master_replace(doc, f"{{{{SD{d}P{p}}}}}", v['subj'])
                    master_replace(doc, f"{{{{TD{d}P{p}}}}}", f"{v['teacher']}{status_tag}")
                buf = BytesIO(); doc.save(buf)
                st.download_button(f"💾 儲存 {target_c} 課表", buf.getvalue(), f"{target_c}_第{current_w}週課表.docx")

    # --- Tab 2: 教師課表 ---
    with tab2:
        teachers = st.session_state.ordered_teachers
        if teachers:
            curr_t = st.session_state.get('sel_teacher', teachers[0])
            if curr_t not in teachers: curr_t = teachers[0]
            
            t_col1, t_col2, t_col3 = st.columns([1, 2, 1])
            if t_col1.button("⬅️ 前一位", key="prev_t"):
                st.session_state.sel_teacher = teachers[(teachers.index(curr_t) - 1) % len(teachers)]; st.rerun()
            if t_col3.button("下一位 ➡️", key="next_t"):
                st.session_state.sel_teacher = teachers[(teachers.index(curr_t) + 1) % len(teachers)]; st.rerun()
            with t_col2: 
                st.session_state.sel_teacher = st.selectbox("跳轉教師", teachers, index=teachers.index(curr_t), key="sb_t")

            target_t = st.session_state.sel_teacher
            base = int(st.session_state.base_hours.get(target_t, 0))
            total = current_total_counts.get(target_t, 0)
            
            m1, m2, m3 = st.columns(3)
            m1.metric("基本應授時數", f"{base} 節")
            m2.metric(f"第 {current_w} 週實際總時數", f"{total} 節")
            m3.metric("本週兼代課/超時數", f"{total-base} 節")
            
            t_prev = []
            for p in range(1, 9):
                row = {"節次": f"第 {p} 節"}
                for d_idx, d_name in enumerate(["一", "二", "三", "四", "五"]):
                    d = d_idx + 1
                    info = st.session_state.schedule_by_week[current_w]["teacher_data"].get(target_t, {}).get((d,p))
                    if info and info["subj"] != "":
                        status_tag = f" [{info['status']}]" if info['status'] != "正常" else ""
                        row[f"週{d_name}"] = f"{info['class']} {info['subj']}{status_tag}"
                    else:
                        row[f"週{d_name}"] = ""
                t_prev.append(row)
            st.table(pd.DataFrame(t_prev))

    # --- Tab 3: 代調課作業核心邏輯 ---
    with tab3:
        st.subheader("🔄 臨時代課與跨課調課處理面板")
        
        op_type = st.radio("請選擇操作項目", ["1. 辦理臨時請假代課（單堂異動）", "2. 辦理跨課課堂對調（雙向調課）"])
        st.divider()
        
        # 提取日期與時間對照表
        w_days_info = st.session_state.weeks_db[current_w]["days"]
        w_date_to_num = st.session_state.weeks_db[current_w]["date_to_day_num"]
        
        if op_type == "1. 辦理臨時請假代課（單堂異動）":
            st.caption("說明：選取請假老師與受更動的課堂，指派代課老師。系統將自動修改當週主畫面的課表與時數統計。")
            with st.form("substitute_form"):
                sc1, sc2, sc3 = st.columns(3)
                with sc1:
                    leave_date = st.date_input("請假課程日期", min_value=st.session_state.weeks_db[1]["start"], max_value=st.session_state.weeks_db[20]["end"])
                with sc2:
                    leave_period = st.selectbox("請假節次", [f"第 {i} 節" for i in range(1, 9)])
                    p_num = int(re.search(r'\d+', leave_period).group())
                with sc3:
                    target_class_for_sub = st.selectbox("發生班級", classes)
                
                # 找出原任課老師是誰
                target_w_num = None
                for w_idx, w_data in st.session_state.weeks_db.items():
                    if leave_date.strftime("%Y-%m-%d") in w_data["date_to_day_num"]:
                        target_w_num = w_idx
                        break
                
                sub_teacher_candidate = st.selectbox("指派代課教師 (B老師)", st.session_state.ordered_teachers)
                
                btn_sub = st.form_submit_button("🔥 執行確認代課變更並產生代課單")
                
                if btn_sub:
                    if not target_w_num:
                        st.error("❌ 選擇的日期不在本學期20週範圍內！")
                    else:
                        d_num, d_name = st.session_state.weeks_db[target_w_num]["date_to_day_num"][leave_date.strftime("%Y-%m-%d")]
                        
                        # 抓出原本的課務資訊
                        orig_info = st.session_state.schedule_by_week[target_w_num]["class_data"][target_class_for_sub].get((d_num, p_num))
                        
                        if not orig_info or orig_info["subj"] == "":
                            st.warning(f"⚠️ 警告：{leave_date} {leave_period} 該班級原本好像沒有排課，將直接強制填入代課資訊。")
                            orig_subj = "臨時課務"
                            orig_teacher = "無"
                        else:
                            orig_subj = orig_info["subj"]
                            orig_teacher = orig_info["teacher"]
                        
                        # 執行方案二：連動修正內部狀態
                        # 1. 修正班級課表裡面的老師與狀態
                        st.session_state.schedule_by_week[target_w_num]["class_data"][target_class_for_sub][(d_num, p_num)] = {
                            "subj": orig_subj, "teacher": sub_teacher_candidate, "status": "代課", "note": f"原{orig_teacher}請假"
                        }
                        # 2. 修正原任課老師課表（將其課務移除或標記請假）
                        if orig_teacher in st.session_state.schedule_by_week[target_w_num]["teacher_data"]:
                            if (d_num, p_num) in st.session_state.schedule_by_week[target_w_num]["teacher_data"][orig_teacher]:
                                st.session_state.schedule_by_week[target_w_num]["teacher_data"][orig_teacher][(d_num, p_num)]["status"] = "請假"
                                st.session_state.schedule_by_week[target_w_num]["teacher_data"][orig_teacher][(d_num, p_num)]["subj"] = "" # 時數扣除
                        
                        # 3. 修正新代課老師課表
                        if sub_teacher_candidate not in st.session_state.schedule_by_week[target_w_num]["teacher_data"]:
                            st.session_state.schedule_by_week[target_w_num]["teacher_data"][sub_teacher_candidate] = {}
                        st.session_state.schedule_by_week[target_w_num]["teacher_data"][sub_teacher_candidate][(d_num, p_num)] = {
                            "subj": orig_subj, "class": target_class_for_sub, "status": "代課", "note": f"代{orig_teacher}"
                        }
                        
                        st.success(f"✅ 已成功將 第{target_w_num}週 週{d_name} {leave_period} {target_class_for_sub} 變更為 {sub_teacher_candidate} 老師代課！主課表已連動更正。")
                        
                        # 同步自動生成 Word 通知書 (套用基礎取代逻辑)
                        sub_doc_bytes = load_default_template("代課樣板.docx")
                        if sub_doc_bytes:
                            doc = Document(BytesIO(sub_doc_bytes))
                            master_replace(doc, "{{DATE}}", leave_date.strftime("%Y-%m-%d"))
                            master_replace(doc, "{{CLASS}}", target_class_for_sub)
                            master_replace(doc, "{{PERIOD}}", leave_period)
                            master_replace(doc, "{{SUBJECT}}", orig_subj)
                            master_replace(doc, "{{LEAVE_TEACHER}}", orig_teacher)
                            master_replace(doc, "{{SUB_TEACHER}}", sub_teacher_candidate)
                            
                            buf = BytesIO(); doc.save(buf)
                            st.download_button("📥 點此下載【代課通知單.docx】", buf.getvalue(), f"{leave_date.strftime('%m%d')}_{target_class_for_sub}_代課通知單.docx")
                        else:
                            st.info("💡 提示：如需下載通知單，請先在後台放入『代課樣板.docx』檔案。")
                            
        elif op_type == "2. 辦理跨課課堂對調（雙向調課）":
            st.caption("說明：設定兩門發生在同班級（或不同班級）的課務進行互相對調（情境 B）。對調後，兩堂課的時間、日期、星期會互換，且雙方老師的總時數維持不變。")
            
            with st.form("exchange_form"):
                st.markdown("##### 📍 第一堂課（課務 A）")
                exa1, exa2, exa3 = st.columns(3)
                with exa1: date_a = st.date_input("日期 (課務 A)", key="da")
                with exa2: period_a = st.selectbox("節次 (課務 A)", [f"第 {i} 節" for i in range(1, 9)], key="pa")
                with exa3: class_a = st.selectbox("班級 (課務 A)", classes, key="ca")
                
                st.markdown("##### 📍 第二堂課（課務 B）")
                exb1, exb2, exb3 = st.columns(3)
                with exb1: date_b = st.date_input("日期 (課務 B)", key="db")
                with exb2: period_b = st.selectbox("節次 (課務 B)", [f"第 {i} 節" for i in range(1, 9)], key="pb")
                with exb3: class_b = st.selectbox("班級 (課務 B)", classes, key="cb")
                
                btn_exchange = st.form_submit_button("🔄 執行確認調課對調並產生調課單")
                
                if btn_exchange:
                    # 找出各自所屬週次
                    w_a, w_b = None, None
                    for w_idx, w_data in st.session_state.weeks_db.items():
                        if date_a.strftime("%Y-%m-%d") in w_data["date_to_day_num"]: w_a = w_idx
                        if date_b.strftime("%Y-%m-%d") in w_data["date_to_day_num"]: w_b = w_idx
                    
                    if not w_a or not w_b:
                        st.error("❌ 選擇的日期有部分超出學期週次範圍！")
                    else:
                        d_num_a, d_name_a = st.session_state.weeks_db[w_a]["date_to_day_num"][date_a.strftime("%Y-%m-%d")]
                        d_num_b, d_name_b = st.session_state.weeks_db[w_b]["date_to_day_num"][date_b.strftime("%Y-%m-%d")]
                        p_num_a = int(re.search(r'\d+', period_a).group())
                        p_num_b = int(re.search(r'\d+', period_b).group())
                        
                        # 讀取原本兩格的資料
                        info_a = st.session_state.schedule_by_week[w_a]["class_data"][class_a].get((d_num_a, p_num_a), {"subj":"","teacher":"","status":"正常"}).copy()
                        info_b = st.session_state.schedule_by_week[w_b]["class_data"][class_b].get((d_num_b, p_num_b), {"subj":"","teacher":"","status":"正常"}).copy()
                        
                        # 交換班級課表內容
                        st.session_state.schedule_by_week[w_a]["class_data"][class_a][(d_num_a, p_num_a)] = {
                            "subj": info_b["subj"], "teacher": info_b["teacher"], "status": "調課", "note": "與B對調"
                        }
                        st.session_state.schedule_by_week[w_b]["class_data"][class_b][(d_num_b, p_num_b)] = {
                            "subj": info_a["subj"], "teacher": info_a["teacher"], "status": "調課", "note": "與A對調"
                        }
                        
                        # 連動修改兩位老師的課表節次對調
                        t_a, t_b = info_a["teacher"], info_b["teacher"]
                        
                        # 移除舊老師在該時段的對照，寫入新對照
                        if t_a and t_a in st.session_state.schedule_by_week[w_a]["teacher_data"]:
                            st.session_state.schedule_by_week[w_a]["teacher_data"][t_a][(d_num_a, p_num_a)] = {"subj": info_b["subj"], "class": class_a, "status": "調課"}
                        if t_b and t_b in st.session_state.schedule_by_week[w_b]["teacher_data"]:
                            st.session_state.schedule_by_week[w_b]["teacher_data"][t_b][(d_num_b, p_num_b)] = {"subj": info_a["subj"], "class": class_b, "status": "調課"}
                        
                        st.success(f"🎉 調課對調完成！【{date_a} {class_a}】已與【{date_b} {class_b}】互換。")
                        
                        # 生成調課 Word 樣板
                        ex_doc_bytes = load_default_template("調課樣板.docx")
                        if ex_doc_bytes:
                            doc = Document(BytesIO(ex_doc_bytes))
                            master_replace(doc, "{{ORG_DATE}}", date_a.strftime("%Y-%m-%d"))
                            master_replace(doc, "{{ORG_WEEK}}", d_name_a)
                            master_replace(doc, "{{ORG_PERIOD}}", period_a)
                            master_replace(doc, "{{ORG_CLASS}}", class_a)
                            master_replace(doc, "{{ORG_TEACHER}}", t_a)
                            master_replace(doc, "{{ORG_SUBJECT}}", info_a["subj"])
                            
                            master_replace(doc, "{{NEW_DATE}}", date_b.strftime("%Y-%m-%d"))
                            master_replace(doc, "{{NEW_WEEK}}", d_name_b)
                            master_replace(doc, "{{NEW_PERIOD}}", period_b)
                            master_replace(doc, "{{NEW_CLASS}}", class_b)
                            master_replace(doc, "{{NEW_TEACHER}}", t_b)
                            master_replace(doc, "{{NEW_SUBJECT}}", info_b["subj"])
                            
                            buf = BytesIO(); doc.save(buf)
                            st.download_button("📥 點此下載【調課通知單.docx】", buf.getvalue(), f"{date_a.strftime('%m%d')}_調課通知單.docx")
                        else:
                            st.info("💡 提示：如需下載通知單，請先在後台放入『調課樣板.docx』檔案。")
else:
    st.info("👋 請至側邊欄設定學期起始日、上傳三個基礎資料檔，並點擊「🚀 執行整合」以解鎖動態多週課表系統。")